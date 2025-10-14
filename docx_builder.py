# docx_builder.py
import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from datetime import datetime
from config import BANNER_PATH, OUTPUT_DIR, THEMES

def create_custom_styles(doc, theme_config):
    """Create custom styles for the document"""
    styles = doc.styles
    
    # Title style
    if 'Report Title' not in [s.name for s in styles]:
        title_style = styles.add_style('Report Title', WD_STYLE_TYPE.PARAGRAPH)
        title_format = title_style.paragraph_format
        title_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_format.space_after = Pt(16)
        title_format.space_before = Pt(8)
        
        title_font = title_style.font
        title_font.name = theme_config["font"]
        title_font.size = Pt(18)
        title_font.bold = True
        title_font.color.rgb = RGBColor.from_string(theme_config["heading_color"])
    
    # Section heading style
    if 'Section Heading' not in [s.name for s in styles]:
        heading_style = styles.add_style('Section Heading', WD_STYLE_TYPE.PARAGRAPH)
        heading_format = heading_style.paragraph_format
        heading_format.space_after = Pt(8)
        heading_format.space_before = Pt(16)
        
        heading_font = heading_style.font
        heading_font.name = theme_config["font"]
        heading_font.size = Pt(theme_config["heading_size"])
        heading_font.bold = True
        heading_font.color.rgb = RGBColor.from_string(theme_config["heading_color"])
    
    # Subsection heading style
    if 'Subsection Heading' not in [s.name for s in styles]:
        sub_heading_style = styles.add_style('Subsection Heading', WD_STYLE_TYPE.PARAGRAPH)
        sub_heading_format = sub_heading_style.paragraph_format
        sub_heading_format.space_after = Pt(4)
        sub_heading_format.space_before = Pt(8)
        
        sub_heading_font = sub_heading_style.font
        sub_heading_font.name = theme_config["font"]
        sub_heading_font.size = Pt(12)
        sub_heading_font.bold = True
        sub_heading_font.color.rgb = RGBColor.from_string(theme_config["accent_color"])
    
    # Body text style
    if 'Body Text Custom' not in [s.name for s in styles]:
        body_style = styles.add_style('Body Text Custom', WD_STYLE_TYPE.PARAGRAPH)
        body_format = body_style.paragraph_format
        body_format.space_after = Pt(8)
        body_format.line_spacing = 1.2
        body_format.first_line_indent = Pt(0)
        
        body_font = body_style.font
        body_font.name = theme_config["font"]
        body_font.size = Pt(theme_config["body_size"])

def build_docx(details, report_text, output_name, theme="Academic Blue", photos=None):
    if not os.path.exists(OUTPUT_DIR):
        os.makedirs(OUTPUT_DIR)

    doc = Document()
    
    # Get theme configuration
    theme_config = THEMES[theme]
    
    # Create custom styles
    create_custom_styles(doc, theme_config)
    
    # Add college header if banner exists
    if os.path.exists(BANNER_PATH):
        header_para = doc.add_paragraph()
        header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = header_para.add_run()
        run.add_picture(BANNER_PATH, width=Inches(6.5))
        
    # Add some space
    doc.add_paragraph()
    
    # Document title
    title_para = doc.add_paragraph()
    title_para.style = 'Report Title'
    title_run = title_para.add_run(f"{details.get('event_title', 'Event')} - Report")
    
    # College info
    college_para = doc.add_paragraph()
    college_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    college_run = college_para.add_run(f"{details.get('college_name', '')}\n{details.get('dept_name', '')}")
    college_run.font.name = theme_config["font"]
    college_run.font.size = Pt(12)
    college_run.font.italic = True
    
    # Add some space
    doc.add_paragraph()
    
    # Process the report text with improved formatting
    sections = [s.strip() for s in report_text.split('\n\n') if s.strip()]
    
    for section in sections:
        if not section:
            continue
            
        lines = [line.strip() for line in section.split('\n') if line.strip()]
        if not lines:
            continue
            
        section_title = lines[0]
        
        # Check if it's a main section heading (starts with **number. or just **)
        if section_title.startswith('**') and section_title.endswith('**'):
            # Main section heading
            heading_para = doc.add_paragraph()
            heading_para.style = 'Section Heading'
            heading_text = section_title.replace('**', '').strip()
            heading_para.add_run(heading_text)
            
            # Process content below heading
            if len(lines) > 1:
                for line in lines[1:]:
                    if line.startswith('**') and line.endswith('**') and len(line) < 100:
                        # Subsection heading
                        sub_heading_para = doc.add_paragraph()
                        sub_heading_para.style = 'Subsection Heading'
                        sub_heading_text = line.replace('**', '').strip()
                        sub_heading_para.add_run(sub_heading_text)
                    elif line.startswith('•') or line.startswith('*'):
                        # Bullet point
                        bullet_para = doc.add_paragraph()
                        bullet_para.style = 'List Bullet'  # Use built-in bullet style
                        bullet_text = line.replace('•', '').replace('*', '').strip()
                        
                        # Handle bold text within bullets
                        if '**' in bullet_text:
                            parts = bullet_text.split('**')
                            for i, part in enumerate(parts):
                                run = bullet_para.add_run(part)
                                if i % 2 == 1:  # Odd indices are bold parts
                                    run.bold = True
                                run.font.name = theme_config["font"]
                                run.font.size = Pt(theme_config["body_size"])
                        else:
                            run = bullet_para.add_run(bullet_text)
                            run.font.name = theme_config["font"]
                            run.font.size = Pt(theme_config["body_size"])
                    elif line.startswith('"') and line.endswith('"'):
                        # Quote
                        quote_para = doc.add_paragraph()
                        quote_para.style = 'Body Text Custom'
                        quote_run = quote_para.add_run(line)
                        quote_run.font.italic = True
                        quote_run.font.name = theme_config["font"]
                        quote_run.font.size = Pt(theme_config["body_size"])
                    else:
                        # Regular paragraph
                        if line.strip():
                            content_para = doc.add_paragraph()
                            content_para.style = 'Body Text Custom'
                            content_para.add_run(line)
        else:
            # Regular content without heading
            if section.strip():
                para = doc.add_paragraph()
                para.style = 'Body Text Custom'
                para.add_run(section)
    
    # Add participant list if provided and not already in AI report
    if details.get("participants") and details["participants"].strip():
        if "participants" not in report_text.lower():
            participants_heading = doc.add_paragraph()
            participants_heading.style = 'Section Heading'
            participants_heading.add_run("Additional Participants")
            
            for participant in details["participants"].split(','):
                if participant.strip():
                    p_para = doc.add_paragraph()
                    p_para.style = 'List Bullet'
                    p_run = p_para.add_run(participant.strip())
                    p_run.font.name = theme_config["font"]
                    p_run.font.size = Pt(theme_config["body_size"])
    
    # Insert images/photos if provided
    if photos:
        photos_heading = doc.add_paragraph()
        photos_heading.style = 'Section Heading'
        photos_heading.add_run("Event Photographs")
        
        for i, photo in enumerate(photos):
            if os.path.exists(photo):
                photo_para = doc.add_paragraph()
                photo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                photo_para.add_run().add_picture(photo, width=Inches(4.5))
                
                # Add caption
                caption_para = doc.add_paragraph()
                caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                caption_run = caption_para.add_run(f"Figure {i+1}: Event Photo")
                caption_run.font.name = theme_config["font"]
                caption_run.font.size = Pt(10)
                caption_run.font.italic = True
    
    # Footer with generation date
    doc.add_paragraph()
    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer_run = footer_para.add_run(f"Report generated on: {datetime.now().strftime('%d %B %Y')}")
    footer_run.font.name = theme_config["font"]
    footer_run.font.size = Pt(9)
    footer_run.font.italic = True
    footer_run.font.color.rgb = RGBColor.from_string("666666")
    
    # Use the provided output_name directly (could be full path or just filename)
    if os.path.dirname(output_name):
        # If output_name contains a directory path, use it directly
        doc.save(output_name)
        return output_name
    else:
        # If output_name is just a filename, save to OUTPUT_DIR
        out_path = os.path.join(OUTPUT_DIR, output_name)
        doc.save(out_path)
        return out_path
